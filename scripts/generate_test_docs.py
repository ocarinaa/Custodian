#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Test Document Generator for Enhanced Document Sorter
Creates various document types for comprehensive system testing.
"""

import os
import sys
from pathlib import Path
from datetime import datetime, timedelta
import random
import json
from typing import List, Dict, Any

class TestDocumentGenerator:
    """Generates realistic test documents for validation."""
    
    def __init__(self, output_dir: str = "test_documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Test scenarios
        self.companies = [
            "ACME Corporation", "TechFlow Ltd", "Global Solutions Inc",
            "Metro Services", "Digital Works", "Prime Industries",
            "Zenith Company", "Alpha Systems", "Beta Dynamics"
        ]
        
        self.people = [
            "John Smith", "Maria Garcia", "Ahmed Hassan", "Li Wei",
            "Emma Johnson", "Carlos Rodriguez", "Priya Patel", "Hans Mueller"
        ]
        
        self.categories = {
            "Invoice": ["INV-{}", "Invoice #{}", "Bill #{}", "Statement #{}"],
            "Contract": ["Contract #{}", "Agreement #{}", "Service Contract {}"],
            "Receipt": ["Receipt #{}", "Purchase Receipt", "Sales Receipt"],
            "Certificate": ["Certificate #{}", "Certification of {}", "Award Certificate"],
            "Letter": ["Letter #{}", "Official Letter", "Business Correspondence"],
            "Report": ["Report #{}", "Analysis Report", "Monthly Report {}"]
        }
    
    def generate_pdf_documents(self) -> List[str]:
        """Generate PDF test documents."""
        generated_files = []
        
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        except ImportError:
            print("Warning: reportlab not available, creating text-based PDFs")
            return self._generate_text_based_pdfs()
        
        for i in range(10):
            category = random.choice(list(self.categories.keys()))
            company = random.choice(self.companies)
            date = datetime.now() - timedelta(days=random.randint(1, 365))
            
            filename = f"test_{category.lower()}_{i+1:03d}.pdf"
            filepath = self.output_dir / filename
            
            # Create PDF
            doc = SimpleDocTemplate(str(filepath), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title = f"{category} Document"
            story.append(Paragraph(title, styles['Title']))
            story.append(Spacer(1, 20))
            
            # Company info
            story.append(Paragraph(f"Company: {company}", styles['Normal']))
            story.append(Paragraph(f"Date: {date.strftime('%Y-%m-%d')}", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Content based on category
            if category == "Invoice":
                story.append(Paragraph(f"Invoice Number: INV-{random.randint(10000, 99999)}", styles['Normal']))
                story.append(Paragraph(f"Amount: ${random.randint(100, 10000)}.{random.randint(10, 99)}", styles['Normal']))
                story.append(Paragraph(f"Due Date: {(date + timedelta(days=30)).strftime('%Y-%m-%d')}", styles['Normal']))
            elif category == "Contract":
                story.append(Paragraph(f"Contract Type: Service Agreement", styles['Normal']))
                story.append(Paragraph(f"Duration: {random.randint(6, 36)} months", styles['Normal']))
                story.append(Paragraph(f"Value: ${random.randint(5000, 100000)}", styles['Normal']))
            elif category == "Receipt":
                story.append(Paragraph(f"Transaction ID: TXN-{random.randint(100000, 999999)}", styles['Normal']))
                story.append(Paragraph(f"Amount Paid: ${random.randint(10, 500)}.{random.randint(10, 99)}", styles['Normal']))
                story.append(Paragraph("Payment Method: Credit Card", styles['Normal']))
            
            # Additional content
            content = f"""
            This is a test document generated for system validation purposes.
            Document contains realistic content to test OCR and AI analysis capabilities.
            
            Additional Information:
            - Reference Number: REF-{random.randint(1000, 9999)}
            - Status: Active
            - Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            """
            story.append(Spacer(1, 20))
            story.append(Paragraph(content, styles['Normal']))
            
            doc.build(story)
            generated_files.append(str(filepath))
        
        return generated_files
    
    def _generate_text_based_pdfs(self) -> List[str]:
        """Generate simple text files as PDF substitutes."""
        generated_files = []
        
        for i in range(10):
            category = random.choice(list(self.categories.keys()))
            company = random.choice(self.companies)
            date = datetime.now() - timedelta(days=random.randint(1, 365))
            
            filename = f"test_{category.lower()}_{i+1:03d}.txt"
            filepath = self.output_dir / filename
            
            content = f"""
{category.upper()} DOCUMENT

Company: {company}
Date: {date.strftime('%Y-%m-%d')}
Document Type: {category}

This is a test document for validation purposes.
Contains realistic content for OCR and AI testing.

Reference: REF-{random.randint(1000, 9999)}
Amount: ${random.randint(100, 5000)}.{random.randint(10, 99)}
Status: Active
            """
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content.strip())
            
            generated_files.append(str(filepath))
        
        return generated_files
    
    def generate_docx_documents(self) -> List[str]:
        """Generate DOCX test documents."""
        generated_files = []
        
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            print("Warning: python-docx not available, creating text files")
            return self._generate_text_based_docx()
        
        for i in range(5):
            category = random.choice(["Contract", "Letter", "Report"])
            company = random.choice(self.companies)
            person = random.choice(self.people)
            date = datetime.now() - timedelta(days=random.randint(1, 180))
            
            filename = f"test_{category.lower()}_{i+1:03d}.docx"
            filepath = self.output_dir / filename
            
            doc = Document()
            
            # Title
            title = doc.add_heading(f'{category} Document', 0)
            
            # Header info
            doc.add_paragraph(f'Company: {company}')
            doc.add_paragraph(f'Contact: {person}')
            doc.add_paragraph(f'Date: {date.strftime("%B %d, %Y")}')
            
            # Content
            if category == "Contract":
                doc.add_heading('Service Agreement', level=1)
                doc.add_paragraph(f'This agreement is between {company} and the service provider.')
                doc.add_paragraph(f'Contract Duration: {random.randint(12, 36)} months')
                doc.add_paragraph(f'Contract Value: ${random.randint(10000, 100000):,}')
                
            elif category == "Letter":
                doc.add_heading('Business Letter', level=1)
                doc.add_paragraph(f'Dear {person},')
                doc.add_paragraph('This is a formal business letter generated for testing purposes.')
                doc.add_paragraph('The content includes realistic business communication.')
                doc.add_paragraph('Best regards,')
                doc.add_paragraph(f'{company} Management')
                
            elif category == "Report":
                doc.add_heading('Business Report', level=1)
                doc.add_paragraph(f'Report prepared by: {person}')
                doc.add_paragraph(f'Report period: {date.strftime("%B %Y")}')
                doc.add_paragraph('Executive Summary:')
                doc.add_paragraph('This report contains analysis and recommendations.')
                doc.add_paragraph(f'Total revenue: ${random.randint(50000, 500000):,}')
            
            doc.save(str(filepath))
            generated_files.append(str(filepath))
        
        return generated_files
    
    def _generate_text_based_docx(self) -> List[str]:
        """Generate text files as DOCX substitutes."""
        generated_files = []
        
        for i in range(5):
            category = random.choice(["Contract", "Letter", "Report"])
            company = random.choice(self.companies)
            person = random.choice(self.people)
            
            filename = f"test_{category.lower()}_{i+1:03d}.txt"
            filepath = self.output_dir / filename
            
            content = f"""
{category.upper()} DOCUMENT

Company: {company}
Contact: {person}
Date: {datetime.now().strftime('%Y-%m-%d')}

This is a {category.lower()} document for testing purposes.
Generated content includes realistic business information.

Reference Number: DOC-{random.randint(1000, 9999)}
Status: Active
            """
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content.strip())
            
            generated_files.append(str(filepath))
        
        return generated_files
    
    def generate_image_documents(self) -> List[str]:
        """Generate image documents with text."""
        generated_files = []
        
        try:
            from PIL import Image, ImageDraw, ImageFont
        except ImportError:
            print("Warning: PIL not available, skipping image generation")
            return []
        
        for i in range(5):
            category = random.choice(["Invoice", "Receipt", "ID Card"])
            company = random.choice(self.companies)
            date = datetime.now() - timedelta(days=random.randint(1, 90))
            
            filename = f"test_{category.lower().replace(' ', '_')}_{i+1:03d}.png"
            filepath = self.output_dir / filename
            
            # Create image
            img = Image.new('RGB', (600, 400), 'white')
            draw = ImageDraw.Draw(img)
            
            try:
                font_large = ImageFont.load_default()
                font_normal = ImageFont.load_default()
            except:
                font_large = None
                font_normal = None
            
            # Draw content
            y_pos = 20
            draw.text((20, y_pos), f"{category.upper()}", fill='black', font=font_large)
            y_pos += 40
            
            draw.text((20, y_pos), f"Company: {company}", fill='black', font=font_normal)
            y_pos += 25
            
            draw.text((20, y_pos), f"Date: {date.strftime('%Y-%m-%d')}", fill='black', font=font_normal)
            y_pos += 25
            
            if category == "Invoice":
                draw.text((20, y_pos), f"Invoice #: INV-{random.randint(10000, 99999)}", fill='black', font=font_normal)
                y_pos += 25
                draw.text((20, y_pos), f"Amount: ${random.randint(100, 2000)}.{random.randint(10, 99)}", fill='black', font=font_normal)
            elif category == "Receipt":
                draw.text((20, y_pos), f"Receipt #: REC-{random.randint(10000, 99999)}", fill='black', font=font_normal)
                y_pos += 25
                draw.text((20, y_pos), f"Total: ${random.randint(10, 500)}.{random.randint(10, 99)}", fill='black', font=font_normal)
            elif category == "ID Card":
                person = random.choice(self.people)
                draw.text((20, y_pos), f"Name: {person}", fill='black', font=font_normal)
                y_pos += 25
                draw.text((20, y_pos), f"ID: {random.randint(100000, 999999)}", fill='black', font=font_normal)
            
            img.save(str(filepath))
            generated_files.append(str(filepath))
        
        return generated_files
    
    def generate_edge_cases(self) -> List[str]:
        """Generate edge case documents for testing error handling."""
        generated_files = []
        
        # Empty file
        empty_file = self.output_dir / "empty_document.txt"
        empty_file.touch()
        generated_files.append(str(empty_file))
        
        # Very large filename
        long_name = "very_long_filename_" + "x" * 100 + ".txt"
        long_file = self.output_dir / long_name
        with open(long_file, 'w') as f:
            f.write("Document with very long filename")
        generated_files.append(str(long_file))
        
        # Special characters in content
        special_file = self.output_dir / "special_chars.txt"
        with open(special_file, 'w', encoding='utf-8') as f:
            f.write("Special chars: àáâãäå çñ €£¥ 中文 العربية русский")
        generated_files.append(str(special_file))
        
        # Corrupted file (invalid extension content)
        corrupt_file = self.output_dir / "corrupted.pdf"
        with open(corrupt_file, 'w') as f:
            f.write("This is not a valid PDF file content")
        generated_files.append(str(corrupt_file))
        
        return generated_files
    
    def generate_manifest(self, all_files: List[str]) -> str:
        """Generate a manifest file with test expectations."""
        manifest = {
            "generated": datetime.now().isoformat(),
            "total_files": len(all_files),
            "test_scenarios": [
                {
                    "scenario": "Normal Processing",
                    "files": [f for f in all_files if "test_" in f and "empty" not in f and "corrupt" not in f],
                    "expected_success_rate": 0.8
                },
                {
                    "scenario": "Edge Cases",
                    "files": [f for f in all_files if any(x in f for x in ["empty", "long", "special", "corrupt"])],
                    "expected_success_rate": 0.3
                },
                {
                    "scenario": "Document Types",
                    "pdf_files": [f for f in all_files if f.endswith(('.pdf', '.txt')) and "test_" in f],
                    "docx_files": [f for f in all_files if "docx" in f or ("txt" in f and any(cat in f for cat in ["contract", "letter", "report"]))],
                    "image_files": [f for f in all_files if f.endswith('.png')]
                }
            ],
            "validation_criteria": {
                "text_extraction": "All readable documents should have text extracted",
                "categorization": "Documents should be categorized correctly based on content",
                "entity_extraction": "Company/person names should be identified",
                "date_parsing": "Dates should be extracted and formatted correctly",
                "error_handling": "Invalid files should be moved to review folder"
            }
        }
        
        manifest_file = self.output_dir / "test_manifest.json"
        with open(manifest_file, 'w') as f:
            json.dump(manifest, f, indent=2)
        
        return str(manifest_file)
    
    def generate_all(self) -> Dict[str, List[str]]:
        """Generate all test documents."""
        print("Generating test documents...")
        
        results = {}
        
        print("  - Creating PDF documents...")
        results['pdf'] = self.generate_pdf_documents()
        
        print("  - Creating DOCX documents...")
        results['docx'] = self.generate_docx_documents()
        
        print("  - Creating image documents...")
        results['images'] = self.generate_image_documents()
        
        print("  - Creating edge case documents...")
        results['edge_cases'] = self.generate_edge_cases()
        
        # Combine all files
        all_files = []
        for file_list in results.values():
            all_files.extend(file_list)
        
        print("  - Creating test manifest...")
        manifest_file = self.generate_manifest(all_files)
        results['manifest'] = [manifest_file]
        
        print(f"\nGeneration complete!")
        print(f"Total files created: {len(all_files)}")
        print(f"Output directory: {self.output_dir}")
        
        # Print summary
        for category, files in results.items():
            print(f"  {category}: {len(files)} files")
        
        return results

def main():
    """Main function."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Generate test documents for Enhanced Document Sorter")
    parser.add_argument("--output", "-o", default="test_documents", 
                       help="Output directory for test documents")
    parser.add_argument("--count", "-c", type=int, default=None,
                       help="Override number of documents to generate")
    
    args = parser.parse_args()
    
    generator = TestDocumentGenerator(args.output)
    results = generator.generate_all()
    
    print("\nTest documents ready for validation!")
    print("To run tests:")
    print(f"1. Set SOURCE_FOLDER={args.output} in your .env file")
    print("2. Run: python test_suite.py")
    print("3. Run: python main_enhanced.py")
    
    return 0

if __name__ == "__main__":
    sys.exit(main())